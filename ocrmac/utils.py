"""Utility functions for input processing."""

import os
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Generator, Tuple, List, Union
from urllib.parse import urlparse

import requests
from PIL import Image

# PDF support
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX support
try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp'}


def is_url(path: str) -> bool:
    """Check if the path is a URL."""
    try:
        result = urlparse(path)
        return result.scheme in ('http', 'https')
    except Exception:
        return False


def is_pdf(path: str) -> bool:
    """Check if the path is a PDF file."""
    return path.lower().endswith('.pdf')


def is_docx(path: str) -> bool:
    """Check if the path is a DOCX file."""
    return path.lower().endswith('.docx')


def is_image(path: str) -> bool:
    """Check if the path is an image file."""
    ext = os.path.splitext(path.lower())[1]
    return ext in IMAGE_EXTENSIONS


def download_file(url: str, timeout: int = 30) -> bytes:
    """Download a file from URL and return its content."""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
    }
    response = requests.get(url, headers=headers, timeout=timeout)
    response.raise_for_status()
    return response.content


def download_image(url: str, timeout: int = 30) -> Image.Image:
    """Download an image from URL and return as PIL Image."""
    content = download_file(url, timeout)
    return Image.open(BytesIO(content))


def pdf_to_images(pdf_path: str, dpi: int = 150) -> Generator[Tuple[int, Image.Image], None, None]:
    """
    Convert PDF pages to PIL Images.

    Args:
        pdf_path: Path to the PDF file
        dpi: Resolution for rendering (default 150)

    Yields:
        Tuple of (page_number, PIL.Image)
    """
    if not PDF_AVAILABLE:
        raise ImportError("PyMuPDF is not installed. Please install it with: pip install PyMuPDF")

    doc = fitz.open(pdf_path)
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Render page to pixmap
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image
            if pix.alpha:
                img = Image.frombytes("RGBA", [pix.width, pix.height], pix.samples)
            else:
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            yield page_num + 1, img
    finally:
        doc.close()


def pdf_from_bytes_to_images(pdf_bytes: bytes, dpi: int = 150) -> Generator[Tuple[int, Image.Image], None, None]:
    """
    Convert PDF bytes to PIL Images.

    Args:
        pdf_bytes: PDF file content as bytes
        dpi: Resolution for rendering (default 150)

    Yields:
        Tuple of (page_number, PIL.Image)
    """
    if not PDF_AVAILABLE:
        raise ImportError("PyMuPDF is not installed. Please install it with: pip install PyMuPDF")

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)

            if pix.alpha:
                img = Image.frombytes("RGBA", [pix.width, pix.height], pix.samples)
            else:
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            yield page_num + 1, img
    finally:
        doc.close()


def extract_docx_content(docx_path: str) -> Tuple[str, List[Tuple[int, Image.Image]]]:
    """
    Extract text and images from a DOCX file.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Tuple of (text_content, list of (index, PIL.Image))
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")

    doc = Document(docx_path)

    # Extract text
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                text_parts.append(" | ".join(row_text))

    text_content = "\n\n".join(text_parts)

    # Extract images
    images = []
    image_index = 0
    try:
        for rel in doc.part.rels.values():
            try:
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    img = Image.open(BytesIO(image_data))
                    image_index += 1
                    images.append((image_index, img))
            except Exception:
                continue
    except Exception:
        pass  # Ignore general relationship errors

    return text_content, images


def extract_docx_from_bytes(docx_bytes: bytes) -> Tuple[str, List[Tuple[int, Image.Image]]]:
    """
    Extract text and images from DOCX bytes.

    Args:
        docx_bytes: DOCX file content as bytes

    Returns:
        Tuple of (text_content, list of (index, PIL.Image))
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")

    doc = Document(BytesIO(docx_bytes))

    # Extract text
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                text_parts.append(" | ".join(row_text))

    text_content = "\n\n".join(text_parts)

    # Extract images
    images = []
    image_index = 0
    try:
        for rel in doc.part.rels.values():
            try:
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    img = Image.open(BytesIO(image_data))
                    image_index += 1
                    images.append((image_index, img))
            except Exception:
                continue
    except Exception:
        pass

    return text_content, images


def get_input_type(path: str) -> str:
    """
    Determine the type of input.

    Returns:
        'url', 'pdf', 'docx', 'image', or 'directory'
    """
    if is_url(path):
        return 'url'

    if os.path.isdir(path):
        return 'directory'

    if is_pdf(path):
        return 'pdf'

    if is_docx(path):
        return 'docx'

    if is_image(path):
        return 'image'

    # Try to detect by content
    return 'unknown'


def get_url_file_type(url: str) -> str:
    """Determine file type from URL."""
    parsed = urlparse(url)
    path = parsed.path.lower()

    if path.endswith('.pdf'):
        return 'pdf'
    elif path.endswith('.docx'):
        return 'docx'
    elif any(path.endswith(ext) for ext in IMAGE_EXTENSIONS):
        return 'image'
    else:
        # Default to image for unknown URLs
        return 'image'


def list_files_in_directory(directory: str, recursive: bool = False) -> List[str]:
    """
    List all supported files in a directory.

    Args:
        directory: Path to directory
        recursive: Whether to search recursively

    Returns:
        List of file paths
    """
    supported_extensions = IMAGE_EXTENSIONS | {'.pdf', '.docx'}
    files = []

    path = Path(directory)
    pattern = '**/*' if recursive else '*'

    for file_path in path.glob(pattern):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            files.append(str(file_path))

    return sorted(files)
